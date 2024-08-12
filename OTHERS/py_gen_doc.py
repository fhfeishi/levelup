import cv2
from docx import Document
from docx.shared import Inches, Cm, Pt 
import os

def test_do_generate_word():
    path = '/home/localproject/qtprojects/MeasureTool/doc/preview.docx'
    save_path = '/home/localproject/qtprojects/MeasureTool/doc/preview_saved.docx'
    pic_path = '/home/localproject/qtprojects/MeasureTool/doc/sample.png'
    replace_dict = {"#abc#":"112233", "#xyz#":"mfm"}

    doc = Document(path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text in replace_dict.keys():
                    cell.text = cell.text.replace(cell.text, replace_dict[cell.text])
                else:
                    print(cell.text)

    run=doc.tables[1].cell(2,2).paragraphs[0].add_run()
    picture =run.add_picture(pic_path)
    picture.height=Cm(7.93)
    picture.width=Cm(10)

    doc.save(save_path)


def do_generate_word(source_path, data_maps, image_path):
    doc = Document(source_path)
    for table in doc.tables:
        table.autofit = True  # 禁用表格的自动调整大小
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # 遍历每一个run，替换占位符
                        for key, value in data_maps.items():
                            run.text = run.text.replace(f'#{key.replace("#","")}#', value)


    source_name = os.path.basename(source_path)
    source_dir = os.path.dirname(source_name)
    save_path = f"{source_dir}/{source_name.replace('.docx','_generated.docx')}"#source_dir + source_name.split('.')[0] + "_gen." + source_name.split('.')[1]
    doc.save(save_path)
    return save_path


